import PyPDF2
import docx2txt
import docx
import json
from openai import OpenAI
import psycopg2
from sqlalchemy import create_engine
import subprocess
import os
import json
import pycountry
from fuzzywuzzy import process


CANDIDATE_TABLE = "candidate_table_test_entire_summary123"
VECTOR_TABLE = "vector_table_test_entire_summary123"
JD_TABLE = "jd_candidate_table_test_entire_summary123"


API_KEY = "xxxxxxxxxxxxxxx"
MODEL = "gpt-4o-mini"
MODEL_4o = "gpt-4o-mini"
EMBEDDINGS_MODEL = "text-embedding-3-large"

# conn_str = f"dbname='{DATABASE}' user='{USER}' host='{HOST}' password='{PASSWORD}'"

# HOST = 'database-1.cvo0sykm8au5.us-east-2.rds.amazonaws.com'
# PORT = '5432'  # default PostgreSQL port is 5432
# DATABASE = 'postgres'
# USER = 'postgres'
# PASSWORD = 'CharanKranthi1!'

HOST = 'localhost' #'localhost' '172.17.0.2' 
#PORT = ''  # default PostgreSQL port is 5432
DATABASE = 'vectordb'
USER = 'postgres'
PASSWORD = 'dev1' #AdminJun!2024
PORT = '5432'
SCORE_LIMIT = 1

# HOST = 'localhost'
# PORT = ''  # default PostgreSQL port is 5432
# DATABASE = 'postgres'
# USER = 'postgres'
# PASSWORD = 'test'


def create_linkedin_validator_prompt(resume_text, linkedin_text):
    prompt_template = f"""

    You are a helpful assistant to read the complete
    candidate resume, candidates linkedin resume and extract relevant infromation which can be used for understanding the candidate skills and other
    details which are helpfull for recruiters

    Below are 2 resumes, one is a candidate resume and the other is his Linkedin resume. Compare the 2 resumes based on skills, location and experience. Make sure to compare all 3 parameter                                                s and generate a response if the 2 resumes are matching with a short description.
    Reason step by step and explain. Write all the reasoning in order like first skills, next experience and then location. Write a details summary if the profiles matches or does not match.
    Candidate Resume :

    {resume_text}

    Linkedin Resume :

    {linkedin_text}

    Your output should be in a json format as described below, including the leading and trailing "json" and "":

        'Compare': string
    """

    return prompt_template



def get_tables(vendor_id):
    CANDIDATE_TABLE = f"candidate_table_{vendor_id}"
    VECTOR_TABLE = f"vector_table_{vendor_id}"
    JD_TABLE = f"jd_candidate_table_{vendor_id}"

    return CANDIDATE_TABLE, VECTOR_TABLE, JD_TABLE


def get_standard_country_name(country_name):
    """Standardize country names using mapping, pycountry, and fuzzy matching"""
    with open("country_mapping.json", "r", encoding="utf-8") as f:
      COUNTRY_MAPPING = json.load(f)
    # Step 1: Check in the predefined mapping
    if country_name in COUNTRY_MAPPING:
        return COUNTRY_MAPPING[country_name]

    # Step 2: Try pycountry lookup
    try:
        country = pycountry.countries.lookup(country_name)
        return country.name
    except LookupError:
        pass

    # Step 3: Fuzzy matching for misspellings and variations
    all_country_names = [c.name for c in pycountry.countries]
    best_match, score = process.extractOne(country_name, all_country_names)

    if score > 80:  # Only accept if confidence is high
        return best_match

    return None  # Return None if no match found


def create_resume_extraction_prompt(text):
    prompt_template = f"""

    Below is a resume. You are a helpfull assistant to read the complete
    resume and extract relevant infromation which can be used for understanding the candidate skills and other
    details which are helpfull for recruiters

    From the resume extract information for the following keywords ['name','location', 'email', 'mobile number',
    'college', 'designation', 'total experience', 'list of companies worked', 'list of skillset',
    'degree', 'list of domains candidate has worked on',hourlyrate, noticeperiod, 'visa status', 'summary']. Identify these details from the entire resume. Make sure to not miss any details                                                .


    Your output should be in a json format as described below, including the leading and trailing "```json" and "```":
    The output json should not include any characters which are not supported by json.
    Return the output in correct json format. Do not include comma after summary. Do not include any new key into the json. Use only below given keys.

        'name': Extract Name of the candidate from the entire resume,
        'location' : Extract the current country of the candidate from the entire resume, if the location is not mentioned get the location from his latest place like his last company,
        'state' : Extract the current state,city of the candidate from the entire resume, if the current state,city is not mentioned get the current state, city from his latest place like his last company. Value should be in state, city format
        'email': Extract email of the candidate from the entire resume,
        'mobileNumber': Extract mobile number of the candidate from the entire resume,
        'college': Extract college name of the candidate from the entire resume,
        'designation': Extract designation of the candidate from the entire resume,
        'totalExperience': Extract total experience in the format 5.1 years,
        'companiesWorked': Extract all the companies that the candidate has worked and create a list,
        'skillset': Extract all the technical skills from the candidate resume,
        'degree': Extract the highest degree of the candidate from the entire resume,
        'domians': Extract all the domains the cadidate has worked on  and create a list,
        'hourlyrate' : Extract the hourly rate of the candicate if mentioned in the  resume,
        'noticeperiod' : Extract the notice period of the candidate like (Immediate, 1 week, 2 weeks, 1 Month, 2 Months, 3 Months) if mentioned in the resume,
        'visastatus': Extract the Visa status of the candidate like (like H1B, H4EAD, OPT EAD, GC, USC, GC EAD) if mentioned in the resume,
        'summary': Analyze the entire resume and write a complete summary. Do not miss any details.


    Here is the resume
    {text}

    If any details are not captured in the json, updated them as 'Not Mentioned'
    """
    return prompt_template


def create_jd_summary_prompt(text):
    prompt_template = f"""

    Below is a Job description. You are a helpfull assistant to read the complete Job description and summarize infromation which can be used for understanding.

    Also write a detailed 500-word summary explaining the candidate's profile with his details. Make sure not to include the candidate's name in the summary. The summary should prioritize t                                                he following elements in order. Use this format to generate summary. Job Designation : . Relevant Experience : . Skills : . Roles and Responsibilities : Extract all the roles and responsibl                                                ities exhibited by the candiade from all his projects. Education : . Certifications: . If any details are not found use Not mentioned.

    Your output should be in a json format as described below, including the leading and trailing "```json" and "```":
    The output json should not include any characters which are not supported by json.
    Return the output in correct json format. Do not include comma after summary. Do not include any new key into the json. Use only below given keys.

        'summary': string

    {text}
    """

    return prompt_template


def get_ratings_prompt(resume, summary):
    prompt_template = f"""

    To assess if the below uploaded resume of the candidate matches the provided job description (JD), we'll evaluate it based on three main criteria: Skills, Experience, and Location. Read the entire resume to get the location. If any location is not mentioned then use 0.

    Below is the Job Description :
    {summary}

    Below is the Resume:
    {resume}
    Compare the Job description against the resume and rate everything in % out of 100.

    Calculate Final Rating using below formula
    Final rating = (Skills rating) * 0.40 + (Experience rating) * 0.40 + (Location) * 0.05 + (HourlyRate) * 0.05 + (Noticeperiod) * 0.05 + (Visastatus) * 0.05

    Do not print the calcualtion in the output. Just print the result.
    Your output should be in a markdown format and strictly use below template.

    ## Resume Assessment

    ### Job Description:
    - **Job Designation**:
    - **Experience**: Compare the cadidate experience with job descrition experience. If the candidate experience is in between Job description experience give 100%. Else give rating accordingly. (Rating out of 100%).
    - **Skills**: Compare the candidate skills and job description skills. Include all the skills from resume and job description. use all the skills mentioned and do not miss any. Check again if you have missed any skills from resume or Job description. append them if not captured.(Rating out of 100%). Write a brief summary of matching and not matching skills.
    - **Location**: Compare the location of the candidate with the location of the job description. If matching give 100% and if it does not match give 0%.
    - **HourlyRate**: Compare the hourlyrate of the candidate with the hourlyrate of the job description. If it matches give 100% else give 0%.
    - **Noticeperiod**: Compare the notice period of the candidate with the notice period of the job description. if it matches give 100% else give 0%.
    - **Visastatus**: Compare the visa status of the candidate with the visa status of the job description.  if it matches give 100% else give 0%.
    ## Final Rating : .

    ### Summary
    - **Summary of observations.***

    """
    return prompt_template



def create_jd_creation_prompt(text):
    prompt_template = f"""
    I want you to act as a Job Description Generator. Your task is to generate a detailed job description based on the format provided below and the information given by the user. If any re                                                quired information is missing or unclear from the user's input, please indicate it as 'ANY' instead of making assumptions.

    Your output should be in a markdown format.

    ---

    Job Title: [Enter Job Title Here]

    Job Location: [Enter Job Location Here]

    Job Type: [Full-time/Part-time/Contract]

    Job Summary:
    [Provide a brief overview of the role and its purpose within the organization.]

    Key Responsibilities:
        [List specific duties and tasks the candidate will be responsible for]
        [List additional responsibilities as necessary]

    Requirements:
        [Education level required]
        [Years of experience]
        [Specific skills or certifications required]
        [Soft skills such as communication, teamwork, etc.]

    Preferred Qualifications:
        [Additional qualifications or experiences that would be beneficial but not mandatory]

    Visa Status:
        [Required Visa Status of the candidate, Keep is as not Applicable when it is not mentioned]

    Compensation and Benefits:
        [Base Salary/ Hourly Rate]
    ---

    Here is the user input:
    {text}
    """
    return prompt_template


def generate_response(prompt_template):
    client = OpenAI(api_key=API_KEY)

    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "user",
                "content": prompt_template,
            }
        ]
    )
    # Extract the content from the response
    return response.choices[0].message.content

def generate_response_4o(prompt_template):
    client = OpenAI(api_key=API_KEY)

    response = client.chat.completions.create(
        model=MODEL_4o,
        temperature=0,
        messages=[
            {
                "role": "user",
                "content": prompt_template,
            }
        ]
    )
    # Extract the content from the response
    return response.choices[0].message.content

def convert_to_json(input_string):
    try:
        # Remove markdown code block markers if present
        cleaned_string = input_string
        if '```' in cleaned_string:
            cleaned_string = cleaned_string.replace('```json', '').replace('```', '').strip()

        # Clean any potential leading/trailing whitespace
        cleaned_string = cleaned_string.strip()

        # Parse the JSON
        data = json.loads(cleaned_string)
        return data
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {str(e)}")
        print(f"Problematic string: {cleaned_string}")
        raise Exception(f"Failed to parse response as JSON: {str(e)}")


def get_embedding(text):
    # Approximate character limit for 8000 tokens
    char_limit = 32000

    # If the text is too long, truncate it
    if len(text) > char_limit:
        text = text[:char_limit]   # text = text.replace("\n", " ")
    client = OpenAI(api_key=API_KEY)
    return client.embeddings.create(input = [text], model=EMBEDDINGS_MODEL, dimensions = 1000).data[0].embedding


def convert_doc_to_docx(input_file, output_dir=None):
    """
    Converts a .doc file to .docx format using LibreOffice.

    :param input_file: Path to the .doc file to be converted.
    :param output_dir: Directory to save the converted file. If None, saves in the same directory as input_file.
    :return: Path to the converted .docx file.
    """
    if not os.path.exists(input_file):
        raise FileNotFoundError(f"The file {input_file} does not exist.")

    if not input_file.endswith('.doc'):
        raise ValueError(f"The file {input_file} is not a .doc file.")

    # Set output directory
    if output_dir is None:
        output_dir = os.path.dirname(input_file)

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Run LibreOffice to convert the file
    try:
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', output_dir, input_file],
            check=True
        )
        converted_file = os.path.join(output_dir, os.path.basename(input_file).replace('.doc', '.docx'))
        if not os.path.exists(converted_file):
            raise FileNotFoundError(f"Conversion failed. {converted_file} not found.")
        return converted_file
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"LibreOffice conversion failed: {e}")
    except Exception as e:
        raise RuntimeError(f"An unexpected error occurred: {e}")


def read_documents(filename):
    if filename.endswith('.pdf'):
        with open(f'{filename}', 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)
            text = ""
            for i in range(num_pages):
                page = pdf_reader.pages[i]
                page_text = page.extract_text()
                if page_text:
                    non_empty_lines = [line for line in page_text.split('\n') if line.strip() != '']
                    text += '\n'.join(non_empty_lines) + '\n'

    if filename.endswith('.docx'):
        doc = docx.Document(f'{filename}')
        # doc = docx.Document(f'./ResumeData/{filename}')
        text = ""
        text = docx2txt.process(filename)
        # Read headers and footers from all sections
        for section in doc.sections:
            # Read headers
            header = section.header
            for paragraph in header.paragraphs:
                text += paragraph.text + "\n"

            # Read footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                text += paragraph.text + "\n"

        # Read main content
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Read tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text += paragraph.text + "\n"
    if filename.endswith('.doc'):
        output_directory = '.'
        converted_file_path = convert_doc_to_docx(filename, output_directory)
        text = read_documents(converted_file_path)



    return text.replace('\0', '')


def get_connection():
    cursor = None
    connection = None
    engine = None
    try:
        # Connect to the PostgreSQL database
        connection = psycopg2.connect(
            host=HOST,
            port=PORT,
            database=DATABASE,
            user=USER,
            password=PASSWORD
        )

        # Create a cursor to perform database operations
        cursor = connection.cursor()
        # record = cursor.fetchone()
        # print("You are connected to - ", record)
        # engine = create_engine(f'postgresql://{USER}:{PASSWORD}@{HOST}:{PORT}/{DATABASE}')
        engine = create_engine(f'postgresql://{USER}:{PASSWORD}@{HOST}/{DATABASE}')
        #engine = create_engine(f'postgresql://{USER}:{PASSWORD}@{HOST}:{PORT}/{DATABASE}')
    except Exception as e:
        print(" Connection failed:", str(e))
    #except (Exception, psycopg2.Error) as error:
        #print("Error while connecting to PostgreSQL", error)

    return cursor, engine, connection

def exec_query(cursor, query, params=None):
    if params:
        cursor.execute(query, params)
    else:
        cursor.execute(query)

def get_records(cursor, table):
    select_query = f"SELECT * FROM {table};"
    cursor.execute(select_query)
    records = cursor.fetchall()

    # Print all records
    for record in records:
        print(record)

